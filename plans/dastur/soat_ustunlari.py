"""Drops the Table 1 hour columns (Ma'ruza / Amaliy / Lab-ya) and section-5
sub-headers for hour types a course has no hours for.

docxtpl substitutes tags but cannot remove table columns or rows outside a
{%tr %} loop, so both are cut out of the rendered document's OOXML instead.
"""

from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tc
from docx.table import Table

from plans.models import FanVariant

JADVAL_INDEKSI = 1
JADVAL_SONI = 3
GRID_USTUNLARI = 10

# Grid column index in Table 1 -> the FanVariant field holding its hours.
SOAT_USTUNLARI: tuple[tuple[int, str], ...] = (
    (4, "maruza_soat"),
    (5, "amaliyot_soat"),
    (6, "laboratoriya_soat"),
)

# Section-5 sub-header row labels, paired with the same fields, in table order.
SOAT_SARLAVHALARI: tuple[tuple[str, str], ...] = (
    ("Ma’ruza (M)", "maruza_soat"),
    ("Amaliy mashg‘ulot (A)", "amaliyot_soat"),
    ("Laboratoriya mashg‘ulot (L)", "laboratoriya_soat"),
)


def soat_ustunlarini_tozala(hujjat: Document, variant: FanVariant) -> None:
    """Removes every hour type this course has no hours for: its column in
    section 1, and its orphan sub-header row in section 5 (that row sits
    outside its {%tr for%} loop, so it survives with an empty loop)."""
    table = _jadval1(hujjat)
    bosh_turlar = [
        maydon for _, maydon in SOAT_USTUNLARI if not getattr(variant, maydon)
    ]
    if len(bosh_turlar) < len(SOAT_USTUNLARI):
        _ustunlarni_ochir(table, bosh_turlar)
    _bosh_sarlavhalarni_ochir(table, variant)


def _ustunlarni_ochir(table: Table, bosh_turlar: list[str]) -> None:
    """Deletes the section-1 grid columns for `bosh_turlar`, then spreads the
    freed width over the surviving hour columns."""
    ochiriladigan = [ustun for ustun, maydon in SOAT_USTUNLARI if maydon in bosh_turlar]

    bosh_kenglik = 0
    for ustun in sorted(ochiriladigan, reverse=True):
        bosh_kenglik += _ustunni_ochir(table, ustun)

    omon_qolganlar = [
        ustun - sum(1 for o in ochiriladigan if o < ustun)
        for ustun, maydon in SOAT_USTUNLARI
        if maydon not in bosh_turlar
    ]
    _ustunlarni_kengaytir(table, omon_qolganlar, bosh_kenglik)
    _kengliklarni_qayta_hisobla(table)


def _bosh_sarlavhalarni_ochir(table: Table, variant: FanVariant) -> None:
    """Removes each section-5 sub-header row whose hour type is zero."""
    for label, maydon in SOAT_SARLAVHALARI:
        if getattr(variant, maydon):
            continue
        tr = _sarlavha_qatorini_top(table, label)
        if tr is not None:
            tr.getparent().remove(tr)


def _sarlavha_qatorini_top(table: Table, label: str):
    """Word splits a paragraph's text across several `w:t` runs (spell-check
    artefacts) — concatenate a row's text before matching, since the label
    is rarely whole in any single run."""
    for tr in table._tbl.findall(qn("w:tr")):
        matn = "".join(t.text or "" for t in tr.iter(qn("w:t")))
        if label in matn:
            return tr
    return None


def _jadval1(hujjat: Document) -> Table:
    if len(hujjat.tables) != JADVAL_SONI:
        raise ValueError(
            f"Kutilgan {JADVAL_SONI} ta jadval, topildi: {len(hujjat.tables)}. "
            "Shablon tuzilishi o‘zgargan bo‘lsa, SOAT_USTUNLARI ham yangilanishi kerak."
        )
    table = hujjat.tables[JADVAL_INDEKSI]
    ustunlar = len(table._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol")))
    if ustunlar != GRID_USTUNLARI:
        raise ValueError(
            f"1-jadvalda {GRID_USTUNLARI} ta grid ustuni kutilgan, topildi: {ustunlar}."
        )
    return table


def _ustunni_ochir(table: Table, ustun: int) -> int:
    """Removes one grid column, returning the width it freed. A cell covering
    the column alone is dropped; one spanning further loses a grid span."""
    grid = table._tbl.find(qn("w:tblGrid"))
    gridcol = grid.findall(qn("w:gridCol"))[ustun]
    bosh_kenglik = int(gridcol.get(qn("w:w")))
    grid.remove(gridcol)

    for tr in table._tbl.findall(qn("w:tr")):
        joriy = 0
        for tc in tr.findall(qn("w:tc")):
            span = _span(tc)
            if joriy <= ustun < joriy + span:
                if span == 1:
                    tr.remove(tc)
                else:
                    tc.tcPr.find(qn("w:gridSpan")).set(qn("w:val"), str(span - 1))
                break
            joriy += span
    return bosh_kenglik


def _ustunlarni_kengaytir(
    table: Table, omon_qolganlar: list[int], bosh_kenglik: int
) -> None:
    """Spreads the freed width evenly over the surviving hour columns so the
    table still spans the full page width; the remainder goes leftmost."""
    if not omon_qolganlar:
        return
    gridcols = table._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol"))
    ulush, qoldiq = divmod(bosh_kenglik, len(omon_qolganlar))
    for tartib, ustun in enumerate(omon_qolganlar):
        gridcol = gridcols[ustun]
        qoshimcha = ulush + (1 if tartib < qoldiq else 0)
        gridcol.set(qn("w:w"), str(int(gridcol.get(qn("w:w"))) + qoshimcha))


def _kengliklarni_qayta_hisobla(table: Table) -> None:
    """Rewrites every cell width from the surviving grid, so no row disagrees
    with the table about where its column boundaries fall."""
    kengliklar = [
        int(gridcol.get(qn("w:w")))
        for gridcol in table._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol"))
    ]
    for tr in table._tbl.findall(qn("w:tr")):
        joriy = 0
        for tc in tr.findall(qn("w:tc")):
            span = _span(tc)
            tcW = tc.tcPr.find(qn("w:tcW"))
            if tcW is not None:
                tcW.set(qn("w:w"), str(sum(kengliklar[joriy : joriy + span])))
            joriy += span


def _span(tc: CT_Tc) -> int:
    gridspan = tc.tcPr.find(qn("w:gridSpan")) if tc.tcPr is not None else None
    return int(gridspan.get(qn("w:val"))) if gridspan is not None else 1
