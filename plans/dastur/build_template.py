"""Builds oquv_dastur.docx from code. Run manually after changing layout:

python -m plans.dastur.build_template
"""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

from plans.dastur.table1_spec import TABLE1_COL_WIDTHS, table1_rows
from plans.dastur.table2_spec import TABLE2_COL_WIDTHS, table2_rows
from plans.dastur.table_helpers import _add_hyperlink, build_table_from_spec
from plans.dastur.template_text import AXBOROT_MANBALARI, IZOH

OUTPUT_PATH = Path(__file__).resolve().parent / "oquv_dastur.docx"
SOURCE_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent.parent
    / "sources"
    / "Kiberxavfsizlik_asoslari_fan_dasturi_ATDT.docx"
)

EXPECTED_TOP_LEVEL_VARS = {
    "university",
    "kafedra",
    "kafedra_council",
    "major",
    "education_form",
    "year",
    "course",
    "academic_years_str",
    "semesters_str",
    "credits_str",
    "weekly_hours_str",
    "plan_number",
    "total_hours",
    "classroom_total",
    "hours",
    "lectures",
    "practicals",
    "labs",
    "seminars",
}


def _set_rfonts(font, name: str) -> None:
    font.name = name
    rpr = font.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def _register_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.size = Pt(14)
    _set_rfonts(normal.font, "Times New Roman")
    normal.paragraph_format.line_spacing = 1.0

    justify = document.styles.add_style("DasturJustify", WD_STYLE_TYPE.PARAGRAPH)
    justify.base_style = normal
    justify.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _load_cover_page(path: Path) -> Document:
    """Loads the reference source .docx and trims its body down to just the
    cover page (title paragraphs + the 1-row approval box), dropping the
    98-row main table and everything after it but keeping the trailing
    sectPr, which carries the page setup (Letter, per the source)."""
    document = Document(path)
    body = document.element.body
    children = list(body)
    main_table_idx = next(
        i
        for i, el in enumerate(children)
        if el.tag == qn("w:tbl") and len(el.findall(qn("w:tr"))) > 5
    )
    for el in children[main_table_idx:]:
        if el.tag != qn("w:sectPr"):
            body.remove(el)
    return document


def _find_paragraph(document: Document, needle: str) -> Paragraph:
    matches = [p for p in document.paragraphs if needle in p.text]
    if not matches:
        raise ValueError(f"Could not find a paragraph containing {needle!r}")
    if len(matches) > 1:
        raise ValueError(
            f"Paragraph text {needle!r} is ambiguous ({len(matches)} hits)"
        )
    return matches[0]


def _replace_run_text(paragraph: Paragraph, old: str, new: str) -> None:
    """Replaces `old` with `new` in the single run that contains it,
    preserving that run's formatting. Raises if no single run holds it, so a
    reference-doc wording change can never silently ship a dead label
    instead of a working Jinja tag."""
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    raise ValueError(f"Could not find {old!r} in a single run of {paragraph.text!r}")


def _consolidate_tag(
    paragraph: Paragraph, start_marker: str, end_marker: str, tag: str
) -> None:
    """Handles the one value Word's autocorrect split across several runs
    (Ta'lim yo'nalishi): puts `tag` whole into the run where `start_marker`
    begins and blanks every run up to and including the one holding
    `end_marker`, so the tag ends up living in exactly one non-empty run —
    required, since docxtpl's tag parser breaks on a tag split across runs.
    """
    runs = paragraph.runs
    start_idx = next(i for i, r in enumerate(runs) if start_marker in r.text)
    end_idx = next(i for i in range(start_idx, len(runs)) if end_marker in runs[i].text)
    prefix = runs[start_idx].text[: runs[start_idx].text.index(start_marker)]
    end_text = runs[end_idx].text
    suffix = end_text[end_text.index(end_marker) + len(end_marker) :]
    if start_idx == end_idx:
        runs[start_idx].text = prefix + tag + suffix
        return
    runs[start_idx].text = prefix + tag
    runs[end_idx].text = suffix
    for i in range(start_idx + 1, end_idx):
        runs[i].text = ""


def _blank_paragraph_after(document: Document, label_text: str) -> None:
    """Clears the reference doc's real author/reviewer name that follows a
    'Tuzuvchi:'/'Taqrizchilar:' label paragraph — kontekst.py has no
    author/reviewer fields, so these stay blank fill-ins as before."""
    paragraphs = document.paragraphs
    idx = next(i for i, p in enumerate(paragraphs) if p.text == label_text)
    for run in paragraphs[idx + 1].runs:
        run.text = ""


def _insert_cover_page_tags(document: Document) -> None:
    _replace_run_text(
        _find_paragraph(document, "KIBERXAVFSIZLIK ASOSLARI"),
        "KIBERXAVFSIZLIK ASOSLARI",
        "{{ course.name|upper }}",
    )
    _replace_run_text(
        _find_paragraph(document, "kunduzgi"), "kunduzgi", "{{ education_form }}"
    )
    _replace_run_text(
        _find_paragraph(document, "Bilim sohasi:"),
        "600000 – Axborot-kommunikatsiya texnologiyalari",
        "{{ major.bilim_sohasi }}",
    )
    _replace_run_text(
        _find_paragraph(document, "Ta’lim sohasi:"),
        "610000 – Axborot-kommunikatsiya texnologiyalari",
        "{{ major.talim_sohasi }}",
    )
    _consolidate_tag(
        _find_paragraph(document, "Ta’lim yo‘nalishi:"),
        "60610",
        "minoti",
        "{{ major.talim_yonalishi }}",
    )
    _replace_run_text(
        _find_paragraph(document, "Namangan – 2026"), "2026", "{{ year }}"
    )

    mazkur = _find_paragraph(document, "Mazkur o‘quv dasturi")
    _replace_run_text(mazkur, "Turan International University", "{{ university }}")
    _replace_run_text(mazkur, "Filologiya", "{{ kafedra }}")
    _replace_run_text(mazkur, "2026-yil", "{{ kafedra_council.year }}-yil")
    _replace_run_text(mazkur, "_" * 19, "{{ kafedra_council.date }}")

    _blank_paragraph_after(document, "Tuzuvchi:")
    _blank_paragraph_after(document, "Taqrizchilar:")


def _apply_reference_hyperlinks(table) -> None:
    """Axborot manbalari rows 14-16 hold www.lex.uz (plain text) then the
    two real hyperlinks; row indices are fixed by table2_rows()'s layout."""
    hyperlinked = [(text, url) for _, text, url in AXBOROT_MANBALARI if url]
    start_row = len(table.rows) - len(hyperlinked)
    for offset, (text, url) in enumerate(hyperlinked):
        paragraph = table.rows[start_row + offset].cells[1].paragraphs[0]
        _add_hyperlink(paragraph, text, url)


def _build_footer_table(document: Document) -> None:
    document.add_paragraph(IZOH, style="DasturJustify")
    table = build_table_from_spec(document, table2_rows(), TABLE2_COL_WIDTHS)
    _apply_reference_hyperlinks(table)


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    run.font.size = Pt(12)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def _build_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(paragraph)


def _add_trailing_paragraphs(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()


def _check_template_variables(path: Path) -> None:
    template = DocxTemplate(path)
    found = template.get_undeclared_template_variables()
    missing = EXPECTED_TOP_LEVEL_VARS - found
    unexpected = found - EXPECTED_TOP_LEVEL_VARS
    if missing:
        raise ValueError(f"Template is missing expected variables: {missing}")
    if unexpected:
        raise ValueError(f"Template has unexpected undeclared variables: {unexpected}")


def build() -> Document:
    document = _load_cover_page(SOURCE_TEMPLATE_PATH)
    _register_styles(document)
    _insert_cover_page_tags(document)
    build_table_from_spec(document, table1_rows(), TABLE1_COL_WIDTHS)
    _build_footer_table(document)
    _add_trailing_paragraphs(document)
    _build_footer(document)
    return document


def main() -> None:
    document = build()
    document.save(OUTPUT_PATH)
    _check_template_variables(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
