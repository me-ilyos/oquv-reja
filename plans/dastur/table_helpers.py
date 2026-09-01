"""Declarative table-building helpers for the o'quv dastur template.

Tables are described as a list of RowSpec (each a list of CellSpec) rather
than built with imperative merge/shade calls per cell, so the 58-row and
17-row tables in build_template.py read as data.
"""

from dataclasses import dataclass
from typing import Literal

from docx.document import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Twips
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

Align = Literal["left", "center", "right", "justify"]

_ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


@dataclass(frozen=True)
class RunSpec:
    text: str
    bold: bool = False
    italic: bool = False


@dataclass(frozen=True)
class CellSpec:
    text: str = ""
    bold: bool = False
    italic: bool = False
    align: Align = "left"
    span: int = 1
    vmerge: Literal["restart", "continue"] | None = None
    shaded: bool = False
    jinja_run: str | None = None
    """A Jinja `{{ }}`/`{%tr %}` tag to write as a single unsplit run,
    instead of `text` — never concatenated from multiple runs, since that
    would break docxtpl's tag parser."""
    runs: tuple[RunSpec, ...] | None = None
    """A label (bold) followed by a Jinja tag (plain) in the same cell —
    e.g. metadata cells like "Fan/modul kodi{{ course.code }}". Each
    RunSpec.text is written as its own run so the tag stays whole."""


@dataclass(frozen=True)
class RowSpec:
    cells: list[CellSpec]


def _shade_cell(cell: _Cell, color: str = "D9D9D9") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def _set_cell_text(
    cell: _Cell, text: str, *, bold: bool, italic: bool, align: Align
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = _ALIGNMENTS[align]
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic


def _add_tag_run(paragraph: Paragraph, tag: str, *, align: Align = "left") -> None:
    """Writes a Jinja `{{ }}`/`{%tr %}` tag as one unsplit run. docxtpl's tag
    parser breaks if the same tag is split across multiple runs, so the tag
    string must always be passed whole to a single add_run() call."""
    paragraph.alignment = _ALIGNMENTS[align]
    paragraph.add_run(tag)


def _add_runs(paragraph: Paragraph, run_specs: tuple[RunSpec, ...], *, align: Align) -> None:
    paragraph.alignment = _ALIGNMENTS[align]
    for run_spec in run_specs:
        run = paragraph.add_run(run_spec.text)
        run.bold = run_spec.bold
        run.italic = run_spec.italic


def _merge_row_cells(table: Table, row_idx: int, spans: list[int]) -> list[_Cell]:
    """Horizontally merges row `row_idx`'s grid cells per `spans`, returns
    the resulting visible cells in order."""
    row_cells = table.rows[row_idx].cells
    merged: list[_Cell] = []
    col = 0
    for span in spans:
        cell = row_cells[col]
        for offset in range(1, span):
            cell = cell.merge(row_cells[col + offset])
        merged.append(cell)
        col += span
    return merged


def _vmerge_cell(cell: _Cell, mode: Literal["restart", "continue"]) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    vmerge = OxmlElement("w:vMerge")
    if mode == "restart":
        vmerge.set(qn("w:val"), "restart")
    tcPr.append(vmerge)


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    """python-docx has no hyperlink API; build <w:hyperlink r:id=...> by
    hand, styled blue/underlined to match a normal Word hyperlink."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def _set_column_widths(table: Table, col_widths: list[int]) -> None:
    grid = table._tbl.find(qn("w:tblGrid"))
    for grid_col, width in zip(grid.findall(qn("w:gridCol")), col_widths, strict=True):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, col_widths, strict=False):
            cell.width = Twips(width)


def build_table_from_spec(
    document: Document, rows: list[RowSpec], col_widths: list[int]
) -> Table:
    table = document.add_table(rows=len(rows), cols=len(col_widths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_column_widths(table, col_widths)

    for row_idx, row_spec in enumerate(rows):
        spans = [cell_spec.span for cell_spec in row_spec.cells]
        cells = _merge_row_cells(table, row_idx, spans)
        for cell, cell_spec in zip(cells, row_spec.cells, strict=True):
            if cell_spec.vmerge is not None:
                _vmerge_cell(cell, cell_spec.vmerge)
            if cell_spec.shaded:
                _shade_cell(cell)
            if cell_spec.vmerge == "continue":
                continue
            _write_cell_content(cell, cell_spec)
    return table


def _write_cell_content(cell: _Cell, cell_spec: CellSpec) -> None:
    paragraph = cell.paragraphs[0]
    if cell_spec.runs is not None:
        _add_runs(paragraph, cell_spec.runs, align=cell_spec.align)
    elif cell_spec.jinja_run is not None:
        _add_tag_run(paragraph, cell_spec.jinja_run, align=cell_spec.align)
    else:
        _set_cell_text(
            cell,
            cell_spec.text,
            bold=cell_spec.bold,
            italic=cell_spec.italic,
            align=cell_spec.align,
        )
