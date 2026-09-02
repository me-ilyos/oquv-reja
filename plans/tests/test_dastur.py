from django.test import TestCase
from docx import Document
from docx.oxml.ns import qn

from accounts.models import Department, Universitet
from plans import dashboard
from plans.dastur.kontekst import dastur_kontekst
from plans.dastur.mavzular import bosh_mavzular
from plans.dastur.service import dastur_egasimi, dastur_render
from plans.models import SoatTuri, Yuklama
from plans.tests.factories import (
    make_fan,
    make_fan_semestr,
    make_guruh,
    make_oqituvchi,
    make_reja,
)


class BoshMavzularTest(TestCase):
    def test_soat_ikkiga_bolinadi(self) -> None:
        mavzular = bosh_mavzular(32, "M")
        self.assertEqual(len(mavzular), 16)
        self.assertEqual(mavzular[0].code, "M1")
        self.assertEqual(mavzular[-1].code, "M16")
        self.assertTrue(all(m.hours == 2 and m.title == "" for m in mavzular))

    def test_toq_soat_qoldiq_qatori(self) -> None:
        mavzular = bosh_mavzular(5, "A")
        self.assertEqual([m.hours for m in mavzular], [2, 2, 1])
        self.assertEqual(mavzular[-1].code, "A3")

    def test_nol_soat_bosh_royxat(self) -> None:
        self.assertEqual(bosh_mavzular(0, "L"), [])
        self.assertEqual(bosh_mavzular(-4, "S"), [])

    def test_ustunlashtirilgan_soat_mavzuda(self) -> None:
        mavzular = bosh_mavzular(10, "", soat_mavzuda=4)
        self.assertEqual([m.hours for m in mavzular], [4, 4, 2])


class DasturEgasimiTest(TestCase):
    def setUp(self) -> None:
        self.yil = dashboard.joriy_akademik_yil()
        self.kafedra = Department.objects.create(nomi="Matematika")
        reja = make_reja(boshlanish_yili=self.yil)
        fan = make_fan(reja, kafedra=self.kafedra)
        self.variant = fan.tanlangan_variant
        self.fs = make_fan_semestr(self.variant, semestr=1)
        self.maruza_egasi = make_oqituvchi(kafedra=self.kafedra)
        self.amaliyot_egasi = make_oqituvchi(kafedra=self.kafedra)
        self.begona = make_oqituvchi(kafedra=self.kafedra)
        Yuklama.objects.create(
            fan_semestr=self.fs, tur=SoatTuri.MARUZA, oqituvchi=self.maruza_egasi
        )
        guruh = make_guruh(reja)
        Yuklama.objects.create(
            fan_semestr=self.fs,
            tur=SoatTuri.AMALIYOT,
            oqituvchi=self.amaliyot_egasi,
            guruh=guruh,
        )

    def test_maruza_egasi_ruxsat_etiladi(self) -> None:
        self.assertTrue(dastur_egasimi(self.variant, self.maruza_egasi))

    def test_amaliyot_egasi_rad_etiladi(self) -> None:
        self.assertFalse(dastur_egasimi(self.variant, self.amaliyot_egasi))

    def test_begona_oqituvchi_rad_etiladi(self) -> None:
        self.assertFalse(dastur_egasimi(self.variant, self.begona))


class DasturEgasimiBoshqaYilTest(TestCase):
    """A curriculum whose year isn't "today's" academic year must not block
    its real lecture owner — ownership isn't tied to the calendar."""

    def test_kelajakdagi_reja_egasi_ham_ruxsat_etiladi(self) -> None:
        kafedra = Department.objects.create(nomi="Fizika")
        kelajak_yil = dashboard.joriy_akademik_yil() + 5
        reja = make_reja(yonalish_kodi="60610199", boshlanish_yili=kelajak_yil)
        fan = make_fan(reja, kafedra=kafedra)
        variant = fan.tanlangan_variant
        fs = make_fan_semestr(variant, semestr=1)
        maruza_egasi = make_oqituvchi(kafedra=kafedra)
        Yuklama.objects.create(
            fan_semestr=fs, tur=SoatTuri.MARUZA, oqituvchi=maruza_egasi
        )
        self.assertTrue(dastur_egasimi(variant, maruza_egasi))


class DasturKontekstTest(TestCase):
    def setUp(self) -> None:
        self.yil = dashboard.joriy_akademik_yil()
        Universitet.objects.create(rasmiy_nomi="Test universiteti")
        self.kafedra = Department.objects.create(
            nomi="Matematika", fakultet="Aniq fanlar fakulteti"
        )
        reja = make_reja(
            boshlanish_yili=self.yil,
            bilim_sohasi_kodi="600000",
            bilim_sohasi_nomi="Aniq fanlar",
        )
        fan = make_fan(reja, raqam="1.05", kafedra=self.kafedra, jami_soat=120)
        self.variant = fan.tanlangan_variant
        self.fs = make_fan_semestr(self.variant, semestr=1)
        self.maruza_egasi = make_oqituvchi(kafedra=self.kafedra)
        self.amaliyot_egasi = make_oqituvchi(kafedra=self.kafedra)
        Yuklama.objects.create(
            fan_semestr=self.fs, tur=SoatTuri.MARUZA, oqituvchi=self.maruza_egasi
        )
        guruh = make_guruh(reja)
        Yuklama.objects.create(
            fan_semestr=self.fs,
            tur=SoatTuri.AMALIYOT,
            oqituvchi=self.amaliyot_egasi,
            guruh=guruh,
        )

    def test_malumot_maydonlari_toldiriladi(self) -> None:
        kontekst = dastur_kontekst(self.variant)
        self.assertEqual(kontekst["university"], "Test universiteti")
        self.assertEqual(kontekst["kafedra"], "Matematika")
        self.assertEqual(kontekst["city"], "Namangan")
        self.assertEqual(kontekst["language"], "O‘zbek")
        self.assertEqual(kontekst["course"]["code"], self.variant.kodi)
        self.assertEqual(kontekst["course"]["name"], self.variant.nomi)
        self.assertEqual(
            kontekst["plan_number"], f"{self.variant.fan.reja.yonalish_kodi} – 1.05"
        )
        self.assertEqual(kontekst["major"]["bilim_sohasi"], "600000 – Aniq fanlar")
        self.assertEqual(kontekst["total_hours"], str(self.variant.fan.jami_soat))
        self.assertEqual(kontekst["classroom_total"], str(self.variant.auditoriya_soat))
        self.assertEqual(kontekst["hours"]["lecture"], str(self.variant.maruza_soat))
        self.assertEqual(kontekst["semesters_str"], "1")
        self.assertEqual(kontekst["credits_str"], str(self.fs.kredit))

    def test_universitet_qatori_yoq_bolsa_standart_nom_ishlatiladi(self) -> None:
        Universitet.objects.all().delete()
        kontekst = dastur_kontekst(self.variant)
        self.assertEqual(kontekst["university"], "Turan International University")

    def test_nol_soat_chiziqcha_bilan_korsatiladi(self) -> None:
        kontekst = dastur_kontekst(self.variant)
        self.assertEqual(kontekst["hours"]["coursework"], "-")
        self.assertEqual(kontekst["hours"]["seminar"], "-")
        self.assertEqual(kontekst["seminars"], [])

    def test_seminar_soat_kontekstga_kiradi(self) -> None:
        self.variant.seminar_soat = 8
        self.variant.save(update_fields=["seminar_soat"])
        kontekst = dastur_kontekst(self.variant)
        self.assertEqual(kontekst["hours"]["seminar"], "8")
        self.assertEqual(len(kontekst["seminars"]), 4)
        self.assertEqual(kontekst["seminars"][0].code, "S1")

    def test_mustaqil_talim_hisoblanadi(self) -> None:
        kontekst = dastur_kontekst(self.variant)
        kutilgan = self.variant.fan.jami_soat - self.variant.auditoriya_soat
        self.assertEqual(kontekst["hours"]["self_study"], str(kutilgan))

    def test_mavzular_soatdan_hisoblanadi(self) -> None:
        kontekst = dastur_kontekst(self.variant)
        self.assertEqual(len(kontekst["lectures"]), self.variant.maruza_soat // 2)
        self.assertEqual(kontekst["practicals"][0].code, "A1")
        self.assertEqual(len(kontekst["labs"]), self.variant.laboratoriya_soat // 2)

    def test_mustaqil_topshiriqlar_4_soatdan_hisoblanadi(self) -> None:
        kontekst = dastur_kontekst(self.variant)
        kutilgan_soat = self.variant.fan.jami_soat - self.variant.auditoriya_soat
        self.assertEqual(len(kontekst["self_study_tasks"]), kutilgan_soat // 4)
        self.assertTrue(all(t.hours == 4 for t in kontekst["self_study_tasks"]))

    def test_oqituvchi_toldiriladigan_maydonlar_bosh(self) -> None:
        kontekst = dastur_kontekst(self.variant)
        self.assertEqual(kontekst["purpose"], "")
        self.assertEqual(kontekst["tasks"], "")
        self.assertEqual(kontekst["prerequisites"], ["", "", ""])
        self.assertEqual(len(kontekst["outcomes"]["professional"]), 8)
        self.assertEqual(len(kontekst["outcomes"]["skills"]), 8)
        self.assertEqual(kontekst["outcomes"]["professional"][0].code, "TN1")

    def test_kop_semestrli_oquv_yillari_takrorlanmaydi(self) -> None:
        make_fan_semestr(self.variant, semestr=2)
        kontekst = dastur_kontekst(self.variant)
        boshlanish = self.variant.fan.reja.boshlanish_yili
        self.assertEqual(
            kontekst["academic_years_str"], f"{boshlanish}/{boshlanish + 1}"
        )

    def test_kengash_maydonlari_bosh_joy(self) -> None:
        kontekst = dastur_kontekst(self.variant)
        self.assertEqual(kontekst["kafedra_council"]["date"], "_____")

    def test_render_hujjat_yaratadi(self) -> None:
        buffer = dastur_render(self.variant)
        self.assertGreater(len(buffer.getvalue()), 0)

    def test_render_hech_qanday_teg_qolmaydi(self) -> None:
        """Every declared variable must be supplied — a leftover {{ }} or
        {%tr %} means one was missed."""
        buffer = dastur_render(self.variant)
        hujjat = Document(buffer)
        self.assertNotRegex(hujjat.element.xml, r"\{\{|\{%")

    def test_render_mavzu_qatorlari_hujjatga_kiradi(self) -> None:
        buffer = dastur_render(self.variant)
        hujjat = Document(buffer)
        matn = "\n".join(
            cell.text
            for table in hujjat.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertIn("M1", matn)
        self.assertIn(f"M{self.variant.maruza_soat // 2}", matn)

    def test_muqova_teglari_toldiriladi(self) -> None:
        buffer = dastur_render(self.variant)
        hujjat = Document(buffer)
        muqova_matni = "\n".join(p.text for p in hujjat.paragraphs)
        self.assertNotRegex(muqova_matni, r"\{\{|\{%")
        self.assertIn(self.variant.nomi.upper(), muqova_matni)
        self.assertIn("600000 – Aniq fanlar", muqova_matni)
        self.assertIn(str(self.yil), muqova_matni)
        self.assertIn(f"{self.kafedra.nomi} kafedrasi", muqova_matni)

    def test_tuzuvchi_va_taqrizchi_bosh_qoldiriladi(self) -> None:
        buffer = dastur_render(self.variant)
        hujjat = Document(buffer)
        paragraphs = hujjat.paragraphs
        muqova_matni = "\n".join(p.text for p in paragraphs)
        self.assertNotIn("Komilov", muqova_matni)
        self.assertNotIn("Xashimov", muqova_matni)

        for label in ("Tuzuvchi:", "Taqrizchilar:"):
            idx = next(i for i, p in enumerate(paragraphs) if p.text == label)
            self.assertEqual(paragraphs[idx + 1].text, "")

    def test_yorliq_va_qiymat_alohida_paragraf(self) -> None:
        """The defect the code-built pipeline had: label and value fused
        onto one line ("Fan/modul turiMajburiy"). They must render as two
        separate paragraphs in the cell, one per line."""
        buffer = dastur_render(self.variant)
        hujjat = Document(buffer)
        matn = "\n".join(
            cell.text
            for table in hujjat.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertIn("Fan/modul turi\nMajburiy", matn)


class DasturSoatUstunlariTest(TestCase):
    """Covers the post-render hour-column/sub-header surgery — it only
    exists in the rendered document, since the raw template always carries
    all three hour columns."""

    def setUp(self) -> None:
        Universitet.objects.create(rasmiy_nomi="Test universiteti")
        kafedra = Department.objects.create(
            nomi="Matematika", fakultet="Aniq fanlar fakulteti"
        )
        reja = make_reja(boshlanish_yili=dashboard.joriy_akademik_yil())
        fan = make_fan(reja, raqam="1.05", kafedra=kafedra, jami_soat=120)
        self.variant = fan.tanlangan_variant
        make_fan_semestr(self.variant, semestr=1)

    def _jadval1(self, variant):
        return Document(dastur_render(variant)).tables[1]

    def _grid_kengliklari(self, jadval) -> list[int]:
        return [
            int(gridcol.get(qn("w:w")))
            for gridcol in jadval._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol"))
        ]

    def _jadval_matni(self, jadval) -> str:
        return "\n".join(cell.text for row in jadval.rows for cell in row.cells)

    def test_barcha_soat_turlari_mavjud_bolsa_hech_narsa_ochirilmaydi(self) -> None:
        self.variant.seminar_soat = 8
        self.variant.save(update_fields=["seminar_soat"])
        jadval = self._jadval1(self.variant)
        matn = self._jadval_matni(jadval)
        for kutilgan in ("Ma’ruza", "Amaliy", "Lab-ya", "Seminar"):
            self.assertIn(kutilgan, matn)
        self.assertEqual(len(self._grid_kengliklari(jadval)), 11)
        # Section 1 values line up with their headers.
        row5 = [c.text for c in jadval.rows[5].cells]
        self.assertIn(str(self.variant.maruza_soat), row5)
        self.assertIn(str(self.variant.amaliyot_soat), row5)
        self.assertIn(str(self.variant.laboratoriya_soat), row5)
        self.assertIn(str(self.variant.seminar_soat), row5)
        # Section 5 carries all four sub-headers, in order.
        for kutilgan in (
            "Ma’ruza (M)",
            "Amaliy mashg‘ulot (A)",
            "Laboratoriya mashg‘ulot (L)",
            "Seminar (S)",
        ):
            self.assertIn(kutilgan, matn)

    def test_laboratoriyasiz_fan_ustuni_va_sarlavhasi_ochiriladi(self) -> None:
        # Factory default seminar_soat is 0, so this drops both Lab-ya and
        # Seminar (2 of the 4 hour columns) -- 11 grid cols down to 9.
        self.variant.laboratoriya_soat = 0
        self.variant.save(update_fields=["laboratoriya_soat"])
        jadval = self._jadval1(self.variant)
        matn = self._jadval_matni(jadval)
        self.assertNotIn("Lab-ya", matn)
        self.assertNotIn("Laboratoriya mashg‘ulot (L)", matn)
        self.assertNotIn("Seminar", matn)
        self.assertIn("Ma’ruza", matn)
        self.assertIn("Amaliy", matn)
        kengliklar = self._grid_kengliklari(jadval)
        self.assertEqual(len(kengliklar), 9)
        self.assertEqual(sum(kengliklar), 9961)

    def test_seminarli_labsiz_fan_ustunlari_togri_ochiriladi(self) -> None:
        """The bug report: InkT.GP 1805 with maruza=32, amaliy=8, seminar=8,
        laboratoriya=0 — everything except Seminar rendered."""
        self.variant.maruza_soat = 32
        self.variant.amaliyot_soat = 8
        self.variant.laboratoriya_soat = 0
        self.variant.seminar_soat = 8
        self.variant.save(
            update_fields=[
                "maruza_soat",
                "amaliyot_soat",
                "laboratoriya_soat",
                "seminar_soat",
            ]
        )
        jadval = self._jadval1(self.variant)
        matn = self._jadval_matni(jadval)
        for kutilgan in ("Ma’ruza", "Amaliy", "Seminar"):
            self.assertIn(kutilgan, matn)
        self.assertNotIn("Lab-ya", matn)
        self.assertNotIn("Laboratoriya mashg‘ulot (L)", matn)
        self.assertIn("Seminar (S)", matn)
        kengliklar = self._grid_kengliklari(jadval)
        self.assertEqual(len(kengliklar), 10)
        self.assertEqual(sum(kengliklar), 9961)

    def test_bitta_soat_turi_qolganda_uchtasi_ochiriladi(self) -> None:
        self.variant.amaliyot_soat = 0
        self.variant.laboratoriya_soat = 0
        self.variant.save(update_fields=["amaliyot_soat", "laboratoriya_soat"])
        jadval = self._jadval1(self.variant)
        matn = self._jadval_matni(jadval)
        self.assertIn("Ma’ruza", matn)
        self.assertNotIn("Amaliy", matn)
        self.assertNotIn("Lab-ya", matn)
        self.assertNotIn("Seminar", matn)
        self.assertEqual(sum(self._grid_kengliklari(jadval)), 9961)

    def test_barcha_soatlar_nol_bolsa_ustunlar_saqlanadi(self) -> None:
        self.variant.maruza_soat = 0
        self.variant.amaliyot_soat = 0
        self.variant.laboratoriya_soat = 0
        self.variant.save(
            update_fields=["maruza_soat", "amaliyot_soat", "laboratoriya_soat"]
        )
        jadval = self._jadval1(self.variant)
        matn = self._jadval_matni(jadval)
        self.assertEqual(len(self._grid_kengliklari(jadval)), 11)
        for kutilgan in ("Ma’ruza", "Amaliy", "Lab-ya"):
            self.assertIn(kutilgan, matn)

    def test_ustun_ochirilgach_vertikal_birlashma_saqlanadi(self) -> None:
        self.variant.laboratoriya_soat = 0
        self.variant.save(update_fields=["laboratoriya_soat"])
        jadval = self._jadval1(self.variant)
        qatorlar = jadval._tbl.findall(qn("w:tr"))

        def vmerge(qator_idx: int):
            return qatorlar[qator_idx].findall(qn("w:tc"))[0].tcPr.find(qn("w:vMerge"))

        self.assertEqual(vmerge(3).get(qn("w:val")), "restart")
        self.assertIsNone(vmerge(4).get(qn("w:val")))


class DasturFormatlashTest(TestCase):
    """Uses the raw template (plans.dastur.service.SABLON_YOLI), not a
    rendered document — docxtpl's {%tr %} loops add/remove rows based on
    context, so row counts are only stable before rendering."""

    def setUp(self) -> None:
        from plans.dastur.service import SABLON_YOLI

        self.hujjat = Document(SABLON_YOLI)

    def test_sahifa_olchami_letter(self) -> None:
        bolim = self.hujjat.sections[0]
        self.assertEqual(bolim.page_width.twips, 12240)
        self.assertEqual(bolim.page_height.twips, 15840)

    def test_jadvallar_soni(self) -> None:
        self.assertEqual(len(self.hujjat.tables), 3)

    def test_yorliq_qiymatdan_alohida_paragraf(self) -> None:
        """Pins the structural fix: label and value are two paragraphs in
        the same cell, not fused into one."""
        cell = self.hujjat.tables[1].rows[2].cells[0]
        self.assertEqual(len(cell.paragraphs), 2)
        self.assertEqual(cell.paragraphs[0].text, "Fan/modul turi")
        self.assertIn("{{ course.module_type }}", cell.paragraphs[1].text)
