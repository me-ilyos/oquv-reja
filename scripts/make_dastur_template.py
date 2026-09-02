"""Build plans/dastur/oquv_dastur.docx from the real reference document.

Run once, and again only if the university reissues the form itself:

    env\\Scripts\\python.exe scripts\\make_dastur_template.py

Opens sources/Kiberxavfsizlik_asoslari_fan_dasturi_ATDT.docx read-only (never
modified) and swaps concrete values for Jinja tags, collapsing repeating rows
into {%tr for %} loops. Keeps every border, merge, and font from the source
byte-for-byte — only text content changes.

After this runs, plans/dastur/oquv_dastur.docx can be opened and tweaked in
Word like any other document (reword a heading, adjust a border); re-run this
script only when the reference document itself changes.
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "sources" / "Kiberxavfsizlik_asoslari_fan_dasturi_ATDT.docx"
TEMPLATE = ROOT / "plans" / "dastur" / "oquv_dastur.docx"


# --------------------------------------------------------------------------
# low-level helpers
# --------------------------------------------------------------------------


def set_para_text(p: Paragraph, text: str) -> None:
    """Replace a paragraph's text, keeping the first run's formatting.

    Word splits text into many runs (spell-check artefacts), so a naive
    search-and-replace misses tags straddling a run boundary. Collapsing to
    one run sidesteps that entirely — verified against the source document,
    where e.g. "60610100 - 1.19" is three separate runs.
    """
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def retext_paragraph(p: Paragraph, replacements: list[tuple[str, str]]) -> None:
    """Collapse a mixed paragraph to one run, applying substring replacements
    to its full text first. For paragraphs where only part of the text
    becomes a tag (e.g. the kafedra-attribution sentence) and the paragraph
    otherwise has no stable per-run split to hook onto.
    """
    text = p.text
    for old, new in replacements:
        if old not in text:
            raise LookupError(f"{old!r} not found in paragraph {p.text!r}")
        text = text.replace(old, new)
    set_para_text(p, text)


def body_para(doc: Document, needle: str) -> Paragraph:
    matches = [p for p in doc.paragraphs if needle in p.text]
    if not matches:
        raise LookupError(f"body paragraph not found: {needle!r}")
    if len(matches) > 1:
        raise LookupError(
            f"paragraph text {needle!r} is ambiguous ({len(matches)} hits)"
        )
    return matches[0]


def cell_para(table: Table, row: int, col: int, idx: int = 0) -> Paragraph:
    """1-based row/col on python-docx's expanded grid; idx is 0-based paragraph."""
    return table.rows[row - 1].cells[col - 1].paragraphs[idx]


def set_cell(table: Table, row: int, col: int, text: str, idx: int = 0) -> None:
    set_para_text(cell_para(table, row, col, idx), text)


def visible_cells(row) -> list[_Cell]:
    """A merged row's `.cells` repeats the same _Cell per grid position; this
    returns each backing cell once, in order."""
    seen: set[int] = set()
    cells = []
    for cell in row.cells:
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        cells.append(cell)
    return cells


def find_row(table: Table, needle: str, *, start: int = 1) -> int:
    """1-based index of the first row (at or after `start`) whose visible
    cells contain `needle`. Anchoring edits on content instead of hardcoded
    row numbers keeps each edit function correct regardless of how many rows
    earlier functions already removed."""
    for i in range(start, len(table.rows) + 1):
        if any(needle in c.text for c in visible_cells(table.rows[i - 1])):
            return i
    raise LookupError(f"no row containing {needle!r} at or after row {start}")


def delete_para(p: Paragraph) -> None:
    p._p.getparent().remove(p._p)


def delete_rows(table: Table, first: int, last: int) -> None:
    """Delete 1-based rows first..last inclusive, bottom-up so indices stay valid."""
    for i in range(last, first - 1, -1):
        tr = table.rows[i - 1]._tr
        tr.getparent().remove(tr)


def _tag_row(row, tag: str) -> None:
    """Blank every cell of a control row and put the tag in the first one."""
    for i, cell in enumerate(visible_cells(row)):
        for j, p in enumerate(cell.paragraphs):
            if j == 0:
                set_para_text(p, tag if i == 0 else "")
            else:
                delete_para(p)


def wrap_rows(table: Table, first: int, last: int, opener: str, closer: str) -> None:
    """Put `opener` in a new row before `first` and `closer` after `last`.

    docxtpl deletes a row whose text is a {%tr %} tag, so the cloned
    formatting of these control rows never reaches rendered output.
    """
    tail = copy.deepcopy(table.rows[last - 1]._tr)
    if last == len(table.rows):
        table._tbl.append(tail)
        closer_row = table.rows[-1]
    else:
        table.rows[last]._tr.addprevious(tail)
        closer_row = table.rows[last]
    _tag_row(closer_row, closer)

    head = copy.deepcopy(table.rows[first - 1]._tr)
    table.rows[first - 1]._tr.addprevious(head)
    _tag_row(table.rows[first - 1], opener)


def _tc_span(tc) -> int:
    pr = tc.find(qn("w:tcPr"))
    gridspan = pr.find(qn("w:gridSpan")) if pr is not None else None
    return int(gridspan.get(qn("w:val"))) if gridspan is not None else 1


def insert_grid_column(
    table: Table,
    after_col: int,
    width: int,
    *,
    split_rows: list[int],
    widen_rows: list[int],
    widen_last_cell_rows: list[int] = (),
) -> None:
    """Inserts a new grid column right after 0-based grid index `after_col`,
    splitting `width` out of that column's own width.

    Only rows named in `split_rows`, `widen_rows`, or `widen_last_cell_rows`
    are touched — every other row is left with its existing gridSpans, which
    silently keep covering the new column (correct for a row whose cell
    already spans past the insertion point, e.g. a full-width section header
    two rows up). `split_rows` get a fresh blank `w:tc` spliced in after the
    cell covering `after_col`; `widen_rows` get that cell's `gridSpan` grown
    by one instead (a header spanning the insertion point widens rather than
    splitting). `widen_last_cell_rows` is for a row whose own cells don't
    straddle `after_col` at all (e.g. its own trailing cell already starts
    after it) but still needs to reach the new final column — its *last*
    cell grows by one, regardless of where `after_col` falls.
    """
    grid = table._tbl.find(qn("w:tblGrid"))
    gridcols = grid.findall(qn("w:gridCol"))
    old_width = int(gridcols[after_col].get(qn("w:w")))
    gridcols[after_col].set(qn("w:w"), str(old_width - width))
    new_gridcol = OxmlElement("w:gridCol")
    new_gridcol.set(qn("w:w"), str(width))
    gridcols[after_col].addnext(new_gridcol)

    targets = set(split_rows) | set(widen_rows)
    for row_idx, tr in enumerate(table._tbl.findall(qn("w:tr")), start=1):
        if row_idx in widen_last_cell_rows:
            tc = tr.findall(qn("w:tc"))[-1]
            span = _tc_span(tc)
            tc.find(qn("w:tcPr")).find(qn("w:gridSpan")).set(qn("w:val"), str(span + 1))
            continue
        if row_idx not in targets:
            continue
        col = 0
        for tc in tr.findall(qn("w:tc")):
            span = _tc_span(tc)
            if col <= after_col < col + span:
                if row_idx in widen_rows:
                    tcPr = tc.find(qn("w:tcPr"))
                    tcPr.find(qn("w:gridSpan")).set(qn("w:val"), str(span + 1))
                else:
                    new_tc = copy.deepcopy(tc)
                    for p in new_tc.findall(qn("w:p")):
                        for run in p.findall(qn("w:r")):
                            p.remove(run)
                    tc.addnext(new_tc)
                break
            col += span

    _recompute_widths(table)


def _recompute_widths(table: Table) -> None:
    """Rewrites every cell's w:tcW from the current tblGrid, so no row
    disagrees with the table about where its column boundaries fall --
    needed after insert_grid_column, since gridSpan changes don't update a
    cell's own stored width."""
    widths = [
        int(gridcol.get(qn("w:w")))
        for gridcol in table._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol"))
    ]
    for tr in table._tbl.findall(qn("w:tr")):
        col = 0
        for tc in tr.findall(qn("w:tc")):
            span = _tc_span(tc)
            tcW = tc.find(qn("w:tcPr")).find(qn("w:tcW"))
            if tcW is not None:
                tcW.set(qn("w:w"), str(sum(widths[col : col + span])))
            col += span


# --------------------------------------------------------------------------
# the edits
# --------------------------------------------------------------------------


def edit_cover(doc: Document) -> None:
    set_para_text(body_para(doc, "KIBERXAVFSIZLIK ASOSLARI"), "{{ course.name|upper }}")
    set_para_text(
        body_para(doc, "kunduzgi ta’lim uchun"), "({{ education_form }} uchun)"
    )
    set_para_text(
        body_para(doc, "Bilim sohasi:"), "Bilim sohasi:\t{{ major.bilim_sohasi }}"
    )
    set_para_text(
        body_para(doc, "Ta’lim sohasi:"), "Ta’lim sohasi:\t{{ major.talim_sohasi }}"
    )
    set_para_text(
        body_para(doc, "Ta’lim yo‘nalishi:"),
        "Ta’lim yo‘nalishi:\t{{ major.talim_yonalishi }}",
    )
    retext_paragraph(
        body_para(doc, "Namangan – 2026"),
        [("Namangan – 2026", "{{ city }} – {{ year }}")],
    )

    kafedra = body_para(doc, "kafedrasi tomonidan taqdim etilgan")
    retext_paragraph(
        kafedra,
        [
            (
                "Turan International University Filologiya kafedrasi tomonidan "
                "taqdim etilgan (kafedraning 2026-yil"
                + "_" * 19
                + "-sonli yig‘ilish bayoni).",
                "{{ university }} {{ kafedra }} kafedrasi tomonidan taqdim etilgan "
                "(kafedraning {{ kafedra_council.year }}-yil "
                "{{ kafedra_council.date }}-sonli yig‘ilish bayoni).",
            )
        ],
    )

    # Approval box: "TASDIQLAYMAN" cell, university name + date fill-ins.
    approval = doc.tables[0].rows[0].cells[1]
    set_para_text(approval.paragraphs[1], "{{ university }}")
    set_para_text(approval.paragraphs[4], "{{ year }}-yil “___” _________")

    # Tuzuvchi/Taqrizchilar: real names in the source, blanked for the teacher.
    set_para_text(body_para(doc, "S.Komilov"), "")
    set_para_text(body_para(doc, "S.Xashimov"), "")


def add_seminar_column(doc: Document) -> None:
    """Splits a new "Seminar" grid column out of the reference document's
    fixed 3-hour-column layout (Ma'ruza/Amaliy/Lab-ya), right after Lab-ya.

    Must run first, against the untouched 98-row table, so the row numbers
    below are stable and correspond 1:1 to the reference document's actual
    layout: row 4 is the "Auditoriya mashg'ulotlari" header (grid 4-6, must
    widen to 4-7 to cover Seminar too); rows 5-6 are the per-column
    sub-header and value rows (each gets its own new blank cell); rows 7-98
    are every full-width row after section 1 (section banners, teacher-fill
    rows, lesson data rows) — they already span the whole table width and
    must widen by one column to keep doing so. Rows 1-3 are left alone: the
    cells that happen to end at grid column 6 there ("Semestr", "O'quv
    rejadagi tartib raqami") are unrelated to the hour columns.
    """
    t = doc.tables[1]
    insert_grid_column(
        t,
        after_col=6,
        width=360,
        split_rows=[5, 6],
        widen_rows=[1, 4, *range(7, len(t.rows) + 1)],
        widen_last_cell_rows=[2, 3],
    )
    set_cell(t, 5, 8, "Seminar")
    set_cell(t, 6, 8, "{{ hours.seminar }}")


def edit_fan_malumotlari(doc: Document) -> None:
    """Section 1, table[1] rows 1-6. Vertically merged block — scalar edits
    only, no rows added or removed. Columns 9+ are one higher than the
    reference document's own numbering, since add_seminar_column() already
    spliced in a new column 8 (Seminar) before this runs."""
    t = doc.tables[1]
    set_cell(t, 2, 1, "{{ course.code }}", idx=1)
    set_cell(t, 2, 4, "{{ academic_years_str }}", idx=1)
    set_cell(t, 2, 5, "{{ semesters_str }}", idx=1)
    set_cell(t, 2, 9, "{{ credits_str }}", idx=1)
    set_cell(t, 3, 1, "{{ course.module_type }}", idx=1)
    set_cell(t, 3, 4, "{{ language }}", idx=1)
    set_cell(t, 3, 5, "{{ plan_number }}", idx=1)
    set_cell(t, 3, 9, "{{ weekly_hours_str }}", idx=1)
    set_cell(t, 4, 1, "{{ course.name }}", idx=1)

    retext_paragraph(
        cell_para(t, 4, 5),
        [("jami – 72 soat, shundan:", "jami – {{ classroom_total }} soat, shundan:")],
    )

    set_cell(t, 6, 4, "{{ total_hours }}")
    set_cell(t, 6, 5, "{{ hours.lecture }}")
    set_cell(t, 6, 6, "{{ hours.practice }}")
    set_cell(t, 6, 7, "{{ hours.lab }}")
    set_cell(t, 6, 9, "{{ hours.self_study }}")
    set_cell(t, 6, 10, "{{ hours.coursework }}")


def edit_mazmuni(doc: Document) -> None:
    """Section 2: purpose and tasks are two paragraphs in one cell."""
    t = doc.tables[1]
    row = find_row(t, "Fanning maqsadi")
    set_cell(t, row, 1, "{{ purpose }}", idx=0)
    set_cell(t, row, 1, "{{ tasks }}", idx=1)


def _collapse_to_one_row(
    t: Table, header_needle: str, first_item_needle: str, group_size: int
) -> int:
    """Deletes rows 2..group_size of a repeating block (found fresh by
    content, so this is safe to call repeatedly without tracking how earlier
    calls shifted row numbers) and returns the surviving first-item row."""
    header = find_row(t, header_needle)
    item1 = find_row(t, first_item_needle, start=header)
    delete_rows(t, item1 + 1, item1 + group_size - 1)
    return item1


def edit_prerequisites(doc: Document) -> None:
    """Section 3: 3 real prerequisites -> loop."""
    t = doc.tables[1]
    item1 = _collapse_to_one_row(t, "Fanni o‘zlashtirish uchun zarur", "(ATD1308)", 3)
    set_cell(t, item1, 1, "{{ loop.index }}.")
    set_cell(t, item1, 2, "{{ item }}")
    wrap_rows(t, item1, item1, "{%tr for item in prerequisites %}", "{%tr endfor %}")


def edit_outcomes(doc: Document) -> None:
    """Section 4: TN1-5 professional, TN6-10 skills -> one data row + loop
    per group. Each anchor is re-found by content after both collapses run,
    since collapsing the (higher-up) professional block shifts every row
    number below it, including the skills block's."""
    t = doc.tables[1]
    _collapse_to_one_row(t, "Ko‘nikmalar:", "TN6", 5)
    _collapse_to_one_row(t, "Kasbiy kompetensiyalar:", "TN1", 5)

    tn1 = find_row(t, "TN1")
    tn6 = find_row(t, "TN6", start=tn1)
    for row in (tn6, tn1):
        set_cell(t, row, 1, "{{ o.code }}")
        set_cell(t, row, 2, "{{ o.text }}")
    wrap_rows(t, tn6, tn6, "{%tr for o in outcomes.skills %}", "{%tr endfor %}")
    wrap_rows(t, tn1, tn1, "{%tr for o in outcomes.professional %}", "{%tr endfor %}")


def edit_topics(doc: Document) -> None:
    """Section 5: M1-12/A1-12/L1-12 -> one data row + loop per lesson type,
    plus a cloned Seminar (S) block the reference document doesn't have.
    Collapses run bottom-up (a block only shifts rows below it), but since
    later collapses (A, then M) sit above earlier ones, every anchor is
    re-found by content once all three collapses are done."""
    t = doc.tables[1]
    _collapse_to_one_row(t, "Laboratoriya mashg‘ulot (L)", "L1", 12)
    _collapse_to_one_row(t, "Amaliy mashg‘ulot (A)", "A1", 12)
    _collapse_to_one_row(t, "Ma’ruza (M)", "M1", 12)

    m1 = find_row(t, "M1")
    a1 = find_row(t, "A1", start=m1)
    l1 = find_row(t, "L1", start=a1)
    a_hdr, l_hdr = a1 - 1, l1 - 1

    for row in (l1, a1, m1):
        set_cell(t, row, 1, "{{ t.code }}")
        set_cell(t, row, 2, "{{ t.title }}")
        set_cell(t, row, 11, "{{ t.hours }}")
    set_cell(t, m1 - 1, 11, "{{ hours.lecture }}")
    set_cell(t, a_hdr, 11, "{{ hours.practice }}")
    set_cell(t, l_hdr, 11, "{{ hours.lab }}")
    wrap_rows(t, l1, l1, "{%tr for t in labs %}", "{%tr endfor %}")
    wrap_rows(t, a1, a1, "{%tr for t in practicals %}", "{%tr endfor %}")
    wrap_rows(t, m1, m1, "{%tr for t in lectures %}", "{%tr endfor %}")

    _add_seminar_block(t)


def _add_seminar_block(t: Table) -> None:
    """Appends a Seminar (S) sub-header + {%tr for%} loop right after the
    Laboratoriya block, same 4-row shape as M/A/L (sub-header, for-open,
    data, endfor -- no if/endif, matching those blocks exactly). Built the
    same way wrap_rows builds a control row: clone an existing row for its
    formatting, then retext every cell. The reference document has no
    seminar section at all, so this whole block is new."""
    l_hdr = find_row(t, "Laboratoriya mashg‘ulot (L)")
    l_endfor = find_row(t, "{%tr endfor %}", start=l_hdr)
    anchor = t.rows[l_endfor - 1]._tr

    for _ in range(4):
        clone = copy.deepcopy(anchor)
        anchor.addnext(clone)
        anchor = clone

    s_hdr, s_for, s_data, s_endfor = range(l_endfor + 1, l_endfor + 5)
    _tag_row(t.rows[s_hdr - 1], "")
    set_cell(t, s_hdr, 2, "Seminar (S)")
    set_cell(t, s_hdr, 11, "{{ hours.seminar }}")
    _tag_row(t.rows[s_for - 1], "{%tr for t in seminars %}")
    set_cell(t, s_data, 1, "{{ t.code }}")
    set_cell(t, s_data, 2, "{{ t.title }}")
    set_cell(t, s_data, 11, "{{ t.hours }}")
    _tag_row(t.rows[s_endfor - 1], "{%tr endfor %}")


def edit_self_study(doc: Document) -> None:
    """Section 6: 27 tasks -> one data row + loop."""
    t = doc.tables[1]
    hdr = find_row(t, "Mustaqil ta’lim topshiriqlari")
    task1 = hdr + 1
    set_cell(t, hdr, 11, "{{ hours.self_study }}")
    delete_rows(t, task1 + 1, len(t.rows))  # tasks 2-27
    set_cell(t, task1, 1, "{{ loop.index }}")
    set_cell(t, task1, 2, "{{ t.title }}")
    set_cell(t, task1, 11, "{{ t.hours }}")
    wrap_rows(t, task1, task1, "{%tr for t in self_study_tasks %}", "{%tr endfor %}")


# --------------------------------------------------------------------------


def main() -> int:
    if not SOURCE.exists():
        print(f"error: {SOURCE} not found", file=sys.stderr)
        return 1

    doc = Document(SOURCE)
    for step in (
        add_seminar_column,
        edit_cover,
        edit_fan_malumotlari,
        edit_mazmuni,
        edit_prerequisites,
        edit_outcomes,
        edit_topics,
        edit_self_study,
    ):
        step(doc)
        print(f"  {step.__name__}")

    TEMPLATE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TEMPLATE)
    print(f"\nwrote {TEMPLATE.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
